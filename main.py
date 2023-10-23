import pandas as pd
import docx

shortListedPS=['SIH1346','SIH1347','SIH1504','SIH1500','SIH1498','SIH1407','SIH1405','SIH1392','SIH1386','SIH1385','SIH1383','SIH1364','SIH1363','SIH1360','SIH1358','SIH1357','SIH1343','SIH1329','SIH1325','SIH1312','SIH1292','SIH1287','SIH1283']
doc = docx.Document() 
problemStatements=pd.read_excel('problemStatements (2).xlsx')
for ps_id in shortListedPS:
    ps=problemStatements[problemStatements['ID']==ps_id]
    id=ps['ID'].values[0]
    title=ps['Title'].values[0]
    desc=ps['Description'].values[0]
    creator=ps['Problem Creater\'s Organization'].values[0]
    tech=ps['Technology Bucket'].values[0]
    doc.add_heading(id)
    doc.add_paragraph("("+tech+")")
    doc.add_heading(title)
    doc.add_paragraph("\n"+desc)
    doc.add_paragraph("By: "+creator)


doc.save('mini_project_PS.docx')